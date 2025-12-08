import os
import json
import sqlite3
import zipfile
import uuid  # === 引入 UUID 库 ===
from io import BytesIO
from datetime import datetime, timedelta
from itertools import groupby
from operator import attrgetter
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, Response, jsonify
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
from PIL import Image
import base64

app = Flask(__name__)
CORS(app)
# 注意：SECRET_KEY 改为新的以区分版本
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'default_key_for_dev') 
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:////data/todo.db'
app.config['UPLOAD_FOLDER'] = '/data/uploads'
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024 
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=30)

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'


# ==========================================
# 核心迁移函数：处理 INTEGER ID 到 UUID ID 的转换
# ==========================================

def get_db_column_info(cursor, table_name):
    """获取表的所有列信息"""
    cursor.execute(f"PRAGMA table_info({table_name})")
    return {info[1]: info[2] for info in cursor.fetchall()} # {column_name: column_type}

def migrate_to_uuid_if_needed(app_context):
    """
    检查 Task 表的主键类型，如果是 INTEGER，则执行复杂迁移。
    """
    db_path = app_context.app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
    if not os.path.exists(db_path):
        print("数据库文件不存在，将直接创建新表。")
        return

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    try:
        task_cols = get_db_column_info(cursor, 'task')
        
        # 1. 检测是否需要迁移：如果 'id' 字段类型是 'INTEGER' 或 'INT'
        if task_cols.get('id', '').upper() not in ['STRING', 'VARCHAR', 'TEXT', 'VARCHAR(36)']:
            print("=== 🚨 检测到旧 INTEGER ID 数据库，开始 UUID 迁移... ===")
            
            # --- 阶段 1: 准备数据和映射 ---
            old_task_id_map = {} # 旧任务 ID (INT) -> 新任务 ID (UUID)
            
            # 备份旧任务数据
            cursor.execute("SELECT id, title, category, priority, content, start_date, due_date, created_at, is_recurring, recurrence_days, completed, completed_at, is_archived, archived_at, user_id FROM task")
            old_tasks = cursor.fetchall()
            
            # 备份旧笔记数据
            cursor.execute("SELECT id, content, images, created_at, task_id FROM note")
            old_notes = cursor.fetchall()

            # --- 阶段 2: 重命名旧表并创建新表 (清空旧结构) ---
            
            # 1. 重命名旧表
            cursor.execute("ALTER TABLE task RENAME TO _task_old")
            cursor.execute("ALTER TABLE note RENAME TO _note_old")
            conn.commit()
            
            # 2. 创建新表 (使用 UUID 架构)
            # 因为我们在 app_context 中，可以直接调用 db.create_all()
            with app_context:
                db.create_all()
            print("新的 UUID 数据库结构已创建。")

            # --- 阶段 3: 导入任务数据 ---
            new_tasks_to_insert = []
            for t in old_tasks:
                old_id = t[0]
                new_uuid = str(uuid.uuid4())
                old_task_id_map[old_id] = new_uuid
                
                # 重新组织数据以匹配新 Task 模型的字段顺序
                # (id, title, category, priority, content, start_date, due_date, created_at, updated_at, is_recurring, recurrence_days, completed, completed_at, is_archived, archived_at, user_id)
                new_tasks_to_insert.append((
                    new_uuid, t[1], t[2], t[3], t[4], t[5], t[6], t[7], t[7], # created_at 和 updated_at 初始值相同
                    t[8], t[9], t[10], t[11], t[12], t[13], t[14]
                ))

            # 批量插入新任务
            task_insert_sql = "INSERT INTO task (id, title, category, priority, content, start_date, due_date, created_at, updated_at, is_recurring, recurrence_days, completed, completed_at, is_archived, archived_at, user_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)"
            cursor.executemany(task_insert_sql, new_tasks_to_insert)
            print(f"任务迁移完成：{len(new_tasks_to_insert)} 条记录。")

            # --- 阶段 4: 导入笔记数据 ---
            new_notes_to_insert = []
            for n in old_notes:
                old_note_id = n[0]
                old_task_id = n[4]
                
                # 查找新的任务 UUID
                new_task_uuid = old_task_id_map.get(old_task_id)
                
                if new_task_uuid:
                    # 重新组织数据以匹配新 Note 模型的字段顺序
                    # (id, content, images, created_at, updated_at, task_id)
                    new_notes_to_insert.append((
                        str(uuid.uuid4()), n[1], n[2], n[3], n[3], # created_at 和 updated_at 初始值相同
                        new_task_uuid
                    ))
            
            # 批量插入新笔记
            note_insert_sql = "INSERT INTO note (id, content, images, created_at, updated_at, task_id) VALUES (?, ?, ?, ?, ?, ?)"
            cursor.executemany(note_insert_sql, new_notes_to_insert)
            print(f"笔记迁移完成：{len(new_notes_to_insert)} 条记录。")

            # --- 阶段 5: 清理 ---
            cursor.execute("DROP TABLE _task_old")
            cursor.execute("DROP TABLE _note_old")
            conn.commit()
            print("旧表已删除，迁移成功。🎉")
            
        else:
            print("数据库结构已是 UUID 格式，跳过迁移。")

    except Exception as e:
        print(f"🚨 数据库迁移失败 (数据未丢失，旧表已备份为 _task_old 和 _note_old): {e}")
        # 确保回滚，防止数据状态不一致
        conn.rollback()

    finally:
        conn.close()

# --- 数据库模型 (UUID重构版) ---

class User(UserMixin, db.Model):
    # 用户表保持 Integer ID 不变
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(150), unique=True, nullable=False)
    password = db.Column(db.String(150), nullable=False)

class Task(db.Model):
    # === 核心变更：ID 改为 UUID 字符串 ===
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    
    title = db.Column(db.String(200), nullable=False)
    category = db.Column(db.String(50), default='其他') 
    priority = db.Column(db.String(20), default='Normal')
    content = db.Column(db.Text)
    start_date = db.Column(db.DateTime)
    due_date = db.Column(db.DateTime)
    
    created_at = db.Column(db.DateTime, default=datetime.now)
    # === 新增：更新时间戳 (用于同步冲突检测) ===
    updated_at = db.Column(db.DateTime, default=datetime.now, onupdate=datetime.now)
    
    is_recurring = db.Column(db.Boolean, default=False)
    recurrence_days = db.Column(db.Integer, default=0)
    completed = db.Column(db.Boolean, default=False)
    completed_at = db.Column(db.DateTime)
    
    is_archived = db.Column(db.Boolean, default=False)
    archived_at = db.Column(db.DateTime)
    
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    notes = db.relationship('Note', backref='task', lazy=True, cascade="all, delete-orphan")
    
    def to_dict(self):
        return {
            'id': self.id, 
            'title': self.title,
            'category': self.category,
            'priority': self.priority,
            'content': self.content,
            'start_date': self.start_date.strftime('%Y-%m-%d %H:%M') if self.start_date else None,
            'due_date': self.due_date.strftime('%Y-%m-%d %H:%M') if self.due_date else None,
            'created_at': self.created_at.strftime('%Y-%m-%d %H:%M') if self.created_at else None,
            'updated_at': self.updated_at.strftime('%Y-%m-%d %H:%M:%S') if self.updated_at else None, 
            'recurrence_days': self.recurrence_days,
            'completed': self.completed,
            'completed_at': self.completed_at.strftime('%Y-%m-%d %H:%M') if self.completed_at else None,
            'is_archived': self.is_archived,
            'archived_at': self.archived_at.strftime('%Y-%m-%d %H:%M') if self.archived_at else None
        }

class Note(db.Model):
    # === 核心变更：ID 改为 UUID 字符串 ===
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    
    content = db.Column(db.Text, nullable=False) 
    images = db.Column(db.Text, default='[]') 
    created_at = db.Column(db.DateTime, default=datetime.now)
    # === 新增：更新时间戳 ===
    updated_at = db.Column(db.DateTime, default=datetime.now, onupdate=datetime.now)
    
    # === 核心变更：外键类型必须与 Task.id 一致 ===
    task_id = db.Column(db.String(36), db.ForeignKey('task.id'), nullable=False)
    
    def get_images(self):
        try: return json.loads(self.images) if self.images else []
        except: return []
    def to_dict(self):
        return {
            'id': self.id, 
            'content': self.content, 
            'images': self.get_images(),
            'created_at': self.created_at.strftime('%Y-%m-%d %H:%M'),
            'updated_at': self.updated_at.strftime('%Y-%m-%d %H:%M:%S') if self.updated_at else None
        }

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

def get_existing_categories(user_id):
    query = db.session.query(Task.category).filter_by(user_id=user_id).distinct()
    categories = [row[0] for row in query if row[0]]
    return sorted(categories)

# --- 辅助函数 ---
def create_thumbnail(image_path):
    try:
        thumb_path = image_path.rsplit('.', 1)[0] + '_thumb.jpg'
        if os.path.exists(thumb_path): return os.path.basename(thumb_path)
        with Image.open(image_path) as img:
            if img.mode in ('RGBA', 'P'): img = img.convert('RGB')
            img.thumbnail((300, 300))
            img.save(thumb_path, "JPEG", quality=70)
        return os.path.basename(thumb_path)
    except Exception as e:
        print(f"缩略图生成失败: {e}")
        return None

def image_to_base64(image_path):
    try:
        if not os.path.exists(image_path): return None
        with open(image_path, "rb") as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
            return f"data:image/jpeg;base64,{encoded_string}"
    except Exception as e: return None

def get_grouped_tasks(user_id, filters):
    query = Task.query.filter_by(user_id=user_id)
    if filters.get('show_archived') == 'true':
        query = query.filter(Task.is_archived == True)
    else:
        query = query.filter((Task.is_archived == False) | (Task.is_archived == None))
    
    if filters.get('q'):
        search = f"%{filters['q']}%"
        query = query.filter((Task.title.like(search)) | (Task.content.like(search)) | (Task.category.like(search)))
    
    if filters.get('category'): query = query.filter(Task.category == filters['category'])

    sort_by = filters.get('sort_by', 'default')
    if sort_by == 'created_desc': query = query.order_by(Task.created_at.desc())
    elif sort_by == 'created_asc': query = query.order_by(Task.created_at.asc())
    elif sort_by == 'completed_desc': query = query.order_by(Task.completed_at.desc())
    elif sort_by == 'due_date': query = query.order_by(Task.due_date.asc())
    else: query = query.order_by(Task.category, Task.completed, Task.due_date)

    tasks = query.all()
    grouped_data = {}
    for task in tasks:
        cat = task.category if task.category and task.category.strip() else '其他'
        if cat not in grouped_data: grouped_data[cat] = []
        grouped_data[cat].append(task)
    return grouped_data

# --- WEB 路由 (UUID 兼容，移除 int: 类型限制) ---

@app.route('/')
@login_required
def dashboard():
    filters = {
        'q': request.args.get('q', ''),
        'category': request.args.get('category', ''),
        'sort_by': request.args.get('sort_by', 'default'),
        'show_archived': request.args.get('show_archived')
    }
    grouped_tasks = get_grouped_tasks(current_user.id, filters)
    categories = get_existing_categories(current_user.id)
    return render_template('dashboard.html', grouped_tasks=grouped_tasks, name=current_user.username, categories=categories, filters=filters, now=datetime.now())

@app.route('/task/<task_id>') # 移除 int:
@login_required
def task_details(task_id):
    task = Task.query.get_or_404(task_id) 
    if task.user_id != current_user.id: return redirect(url_for('dashboard'))
    return render_template('task_details.html', task=task)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        action = request.form.get('action')
        user = User.query.filter_by(username=username).first()
        if action == 'register':
            if user: flash('用户名已存在')
            else:
                new_user = User(username=username, password=generate_password_hash(password, method='scrypt'))
                db.session.add(new_user)
                db.session.commit()
                login_user(new_user)
                return redirect(url_for('dashboard'))
        else: 
            if user and check_password_hash(user.password, password):
                login_user(user)
                return redirect(url_for('dashboard'))
            else: flash('用户名或密码错误')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout(): logout_user(); return redirect(url_for('login'))

@app.route('/add_task', methods=['POST'])
@login_required
def add_task():
    new_task = Task(
        title=request.form.get('title'), 
        category=request.form.get('category') or '其他', 
        priority=request.form.get('priority'), 
        content=request.form.get('content'), 
        user_id=current_user.id,
        start_date=datetime.strptime(request.form.get('start_date'), '%Y-%m-%dT%H:%M') if request.form.get('start_date') else None,
        due_date=datetime.strptime(request.form.get('due_date'), '%Y-%m-%dT%H:%M') if request.form.get('due_date') else None,
        is_recurring=True if request.form.get('recurrence_days') and int(request.form.get('recurrence_days')) > 0 else False,
        recurrence_days=int(request.form.get('recurrence_days') or 0)
    )
    db.session.add(new_task)
    db.session.commit() # updated_at/created_at 自动设置
    return redirect(url_for('dashboard'))

@app.route('/edit_task', methods=['POST'])
@login_required
def edit_task():
    task = Task.query.get_or_404(request.form.get('task_id')) # task_id 是 UUID 字符串
    if task.user_id == current_user.id:
        task.title = request.form.get('title')
        task.category = request.form.get('category') or '其他'
        task.content = request.form.get('content')
        task.priority = request.form.get('priority')
        start = request.form.get('start_date')
        due = request.form.get('due_date')
        task.start_date = datetime.strptime(start, '%Y-%m-%dT%H:%M') if start else None
        task.due_date = datetime.strptime(due, '%Y-%m-%dT%H:%M') if due else None
        task.recurrence_days = int(request.form.get('recurrence_days') or 0)
        task.is_recurring = task.recurrence_days > 0
        db.session.commit() # updated_at 自动更新
    return redirect(url_for('dashboard'))

@app.route('/complete/<id>') # 移除 int:
@login_required
def complete_task(id):
    task = Task.query.get_or_404(id)
    if task.user_id == current_user.id:
        task.completed = not task.completed
        task.completed_at = datetime.now() if task.completed else None
        if task.completed and task.is_recurring and task.recurrence_days > 0:
            next_due = task.due_date + timedelta(days=task.recurrence_days) if task.due_date else datetime.now() + timedelta(days=task.recurrence_days)
            # 新任务使用新的 UUID
            new_task = Task(title=f"{task.title} (循环)", category=task.category, priority=task.priority, content=task.content, user_id=task.user_id, start_date=datetime.now(), due_date=next_due, is_recurring=True, recurrence_days=task.recurrence_days)
            db.session.add(new_task)
            flash('循环任务已生成')
        db.session.commit()
    return redirect(request.referrer or url_for('dashboard'))

@app.route('/delete/<id>') # 移除 int:
@login_required
def delete_task(id):
    task = Task.query.get_or_404(id)
    if task.user_id == current_user.id: db.session.delete(task); db.session.commit()
    return redirect(url_for('dashboard'))

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    from flask import send_from_directory
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/archive/<id>') # 移除 int:
@login_required
def archive_task(id):
    task = Task.query.get_or_404(id)
    if task.user_id == current_user.id:
        task.is_archived = True
        task.archived_at = datetime.now()
        db.session.commit()
        flash('任务已移入归档箱')
    return redirect(request.referrer or url_for('dashboard'))

@app.route('/unarchive/<id>') # 移除 int:
@login_required
def unarchive_task(id):
    task = Task.query.get_or_404(id)
    if task.user_id == current_user.id:
        task.is_archived = False
        task.archived_at = None
        db.session.commit()
        flash('任务已恢复')
    return redirect(request.referrer or url_for('dashboard'))

def create_task_docx(task):
    doc = Document()
    doc.add_heading(task.title, 0)
    p = doc.add_paragraph()
    p.add_run(f"分类: {task.category} | ").bold = True
    p.add_run(f"创建时间: {task.created_at.strftime('%Y-%m-%d')}")
    if task.content: doc.add_paragraph(task.content)
    doc.add_page_break()
    for note in task.notes:
        doc.add_heading(note.created_at.strftime('%Y-%m-%d %H:%M'), level=2)
        doc.add_paragraph(note.content)
        for img in note.get_images():
            img_path = os.path.join(app.config['UPLOAD_FOLDER'], img)
            if os.path.exists(img_path):
                try: doc.add_picture(img_path, width=Inches(5.5))
                except: doc.add_paragraph(f"[图片加载失败: {img}]")
    return doc

@app.route('/download_task/<task_id>') # 移除 int:
@login_required
def download_task(task_id):
    task = Task.query.get_or_404(task_id)
    if task.user_id != current_user.id: return redirect(url_for('dashboard'))
    doc = create_task_docx(task)
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return send_file(f, as_attachment=True, download_name=secure_filename(f"{task.title}.docx"))

@app.route('/batch_action', methods=['POST'])
@login_required
def batch_action():
    task_ids = request.form.getlist('task_ids[]') 
    action_type = request.form.get('action_type')
    if not task_ids: return redirect(request.referrer)
    tasks = Task.query.filter(Task.id.in_(task_ids)).all()
    
    if action_type == 'archive':
        for task in tasks:
            if task.user_id == current_user.id:
                task.is_archived = True
                task.archived_at = datetime.now()
        db.session.commit()
    elif action_type == 'delete':
        for task in tasks:
            if task.user_id == current_user.id: db.session.delete(task)
        db.session.commit()
    elif action_type == 'export':
        memory_file = BytesIO()
        with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
            for task in tasks:
                if task.user_id != current_user.id: continue
                doc = create_task_docx(task)
                docx_io = BytesIO()
                doc.save(docx_io)
                safe_title = secure_filename(task.title) or f"task_{task.id}"
                zf.writestr(f"{safe_title}.docx", docx_io.getvalue())
        memory_file.seek(0)
        return send_file(memory_file, download_name=f"export_{datetime.now().strftime('%Y%m%d')}.zip", as_attachment=True)
    return redirect(request.referrer)

# ==========================================
# API 接口区域：UUID 支持 + 离线同步
# ==========================================

@app.route('/api/tasks', methods=['GET'])
def api_get_tasks():
    auth = request.authorization
    if not auth: return jsonify({'error': 'Auth required'}), 401
    user = User.query.filter_by(username=auth.username).first()
    if not user or not check_password_hash(user.password, auth.password): return jsonify({'error': 'Invalid'}), 401

    show_archived = request.args.get('show_archived', 'false') == 'true'
    sort_by = request.args.get('sort_by', 'default')
    q = request.args.get('q')
    
    query = Task.query.filter_by(user_id=user.id)
    if show_archived: query = query.filter(Task.is_archived == True)
    else: query = query.filter((Task.is_archived == False) | (Task.is_archived == None))
    
    if q:
        search = f"%{q}%"
        query = query.filter((Task.title.like(search)) | (Task.content.like(search)) | (Task.category.like(search)))
        
    if sort_by == 'created_desc': query = query.order_by(Task.created_at.desc())
    elif sort_by == 'created_asc': query = query.order_by(Task.created_at.asc())
    elif sort_by == 'completed_desc': query = query.order_by(Task.completed_at.desc())
    elif sort_by == 'due_date': query = query.order_by(Task.due_date.asc())
    else: query = query.order_by(Task.created_at.desc())

    tasks = query.all()
    data = []
    upload_folder = app.config['UPLOAD_FOLDER']
    for t in tasks:
        item = t.to_dict()
        item['notes'] = []
        for n in t.notes:
            note_dict = n.to_dict()
            images_info = []
            for img in n.get_images():
                full_url = url_for('uploaded_file', filename=img, _external=True)
                thumb_name = img.rsplit('.', 1)[0] + '_thumb.jpg'
                thumb_path = os.path.join(upload_folder, thumb_name)
                if not os.path.exists(thumb_path): create_thumbnail(os.path.join(upload_folder, img))
                base64_str = image_to_base64(thumb_path)
                images_info.append({'original_url': full_url, 'thumb_base64': base64_str, 'filename': img})
            note_dict['images_info'] = images_info
            item['notes'].append(note_dict)
        data.append(item)
    return jsonify({'status': 'success', 'data': data})

# 2. 新增任务 (支持客户端生成 UUID)
@app.route('/api/tasks', methods=['POST'])
def api_create_task():
    auth = request.authorization
    if not auth: return jsonify({'error': 'Auth required'}), 401
    user = User.query.filter_by(username=auth.username).first()
    if not user or not check_password_hash(user.password, auth.password): return jsonify({'error': 'Invalid'}), 401

    data = request.json 
    if not data or 'title' not in data: return jsonify({'error': 'Title is required'}), 400

    start_date, due_date = None, None
    if data.get('start_date'):
        try: start_date = datetime.strptime(data['start_date'], '%Y-%m-%d %H:%M')
        except: pass
    if data.get('due_date'):
        try: due_date = datetime.strptime(data['due_date'], '%Y-%m-%d %H:%M')
        except: pass
        
    task_id = data.get('id', str(uuid.uuid4()))
    
    if Task.query.get(task_id):
        return jsonify({'error': 'Task ID already exists'}), 409

    new_task = Task(
        id=task_id,
        title=data['title'],
        category=data.get('category', '其他'),
        content=data.get('content', ''),
        priority=data.get('priority', 'Normal'),
        start_date=start_date,
        due_date=due_date,
        user_id=user.id,
        created_at=datetime.now()
    )
    db.session.add(new_task)
    db.session.commit()
    
    return jsonify({'status': 'success', 'message': 'Task created', 'id': new_task.id})

# 3. 单个任务操作 (支持 UUID URL)
@app.route('/api/tasks/<task_id>', methods=['GET', 'PUT', 'DELETE']) # 移除 int:
def api_task_action(task_id):
    auth = request.authorization
    if not auth: return jsonify({'error': 'Auth required'}), 401
    user = User.query.filter_by(username=auth.username).first()
    if not user or not check_password_hash(user.password, auth.password): return jsonify({'error': 'Invalid'}), 401

    task = Task.query.get(task_id)
    if not task or task.user_id != user.id: return jsonify({'error': 'Task not found'}), 404

    if request.method == 'GET':
        item = task.to_dict()
        item['notes'] = []
        upload_folder = app.config['UPLOAD_FOLDER']
        for n in task.notes:
            note_dict = n.to_dict()
            images_info = []
            for img in n.get_images():
                full_url = url_for('uploaded_file', filename=img, _external=True)
                thumb_name = img.rsplit('.', 1)[0] + '_thumb.jpg'
                thumb_path = os.path.join(upload_folder, thumb_name)
                if not os.path.exists(thumb_path): create_thumbnail(os.path.join(upload_folder, img))
                base64_str = image_to_base64(thumb_path)
                images_info.append({'original_url': full_url, 'thumb_base64': base64_str, 'filename': img})
            note_dict['images_info'] = images_info
            item['notes'].append(note_dict)
        return jsonify({'status': 'success', 'data': item})

    elif request.method == 'PUT':
        data = request.json
        if 'title' in data: task.title = data['title']
        if 'content' in data: task.content = data['content']
        if 'category' in data: task.category = data['category']
        if 'priority' in data: task.priority = data['priority']
        
        if 'start_date' in data:
            try: task.start_date = datetime.strptime(data['start_date'], '%Y-%m-%d %H:%M') if data['start_date'] else None
            except: pass
        if 'due_date' in data:
            try: task.due_date = datetime.strptime(data['due_date'], '%Y-%m-%d %H:%M') if data['due_date'] else None
            except: pass
            
        if 'completed' in data: 
            task.completed = bool(data['completed'])
            if task.completed and not task.completed_at: task.completed_at = datetime.now()
            elif not task.completed: task.completed_at = None
        
        if 'is_archived' in data:
            task.is_archived = bool(data['is_archived'])
            if task.is_archived and not task.archived_at: task.archived_at = datetime.now()
            elif not task.is_archived: task.archived_at = None

        db.session.commit() # updated_at 自动刷新
        return jsonify({'status': 'success', 'message': 'Task updated'})

    elif request.method == 'DELETE':
        db.session.delete(task)
        db.session.commit()
        return jsonify({'status': 'success', 'message': 'Task deleted'})

# 4. 新增笔记 (支持客户端生成 UUID)
@app.route('/api/notes', methods=['POST'])
def api_add_note():
    auth = request.authorization
    if not auth: return jsonify({'error': 'Auth required'}), 401
    user = User.query.filter_by(username=auth.username).first()
    if not user or not check_password_hash(user.password, auth.password): return jsonify({'error': 'Invalid'}), 401

    task_id = request.form.get('task_id')
    content = request.form.get('content', '')
    
    note_id = request.form.get('id', str(uuid.uuid4()))
    
    if not task_id: return jsonify({'error': 'Task ID required'}), 400
    task = Task.query.get(task_id) 
    if not task or task.user_id != user.id: return jsonify({'error': 'Task not found'}), 404

    saved_images = []
    files = request.files.getlist('images')
    for file in files:
        if file and file.filename:
            filename = secure_filename(f"{datetime.now().timestamp()}_{file.filename}")
            save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(save_path)
            create_thumbnail(save_path)
            saved_images.append(filename)

    new_note = Note(
        id=note_id, 
        content=content, 
        images=json.dumps(saved_images), 
        task_id=task.id
    )
    db.session.add(new_note)
    db.session.commit()
    return jsonify({'status': 'success', 'message': 'Note added', 'note_id': new_note.id})

# 5. 编辑笔记 (支持 UUID Note ID)
@app.route('/api/notes/<note_id>', methods=['PUT']) # 移除 int:
def api_edit_note(note_id):
    auth = request.authorization
    if not auth: return jsonify({'error': 'Auth required'}), 401
    user = User.query.filter_by(username=auth.username).first()
    if not user or not check_password_hash(user.password, auth.password): return jsonify({'error': 'Invalid'}), 401

    note = Note.query.get(note_id)
    if not note or note.task.user_id != user.id: return jsonify({'error': 'Note not found'}), 404

    if 'content' in request.form: note.content = request.form.get('content')
    current_images = note.get_images()
    delete_images = request.form.getlist('delete_images')
    for img in delete_images:
        if img in current_images: current_images.remove(img)
    new_files = request.files.getlist('new_images')
    for file in new_files:
        if file and file.filename:
            filename = secure_filename(f"{datetime.now().timestamp()}_{file.filename}")
            save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(save_path)
            create_thumbnail(save_path)
            current_images.append(filename)
    
    note.images = json.dumps(current_images)
    db.session.commit()
    return jsonify({'status': 'success', 'message': 'Note updated'})

# 6. 删除笔记 (支持 UUID Note ID)
@app.route('/api/notes/<note_id>', methods=['DELETE']) # 移除 int:
def api_delete_note(note_id):
    auth = request.authorization
    if not auth: return jsonify({'error': 'Auth required'}), 401
    user = User.query.filter_by(username=auth.username).first()
    if not user or not check_password_hash(user.password, auth.password): return jsonify({'error': 'Invalid'}), 401

    note = Note.query.get(note_id)
    if not note or note.task.user_id != user.id: return jsonify({'error': 'Not found'}), 404
    
    db.session.delete(note)
    db.session.commit()
    return jsonify({'status': 'success', 'message': 'Note deleted'})

if __name__ == '__main__':
    with app.app_context():
        # --- 启动时执行迁移 ---
        # 1. 检查并迁移数据到 UUID 结构
        migrate_to_uuid_if_needed(app.app_context()) 
        # 2. 如果是新数据库，或迁移成功，确保调用 create_all 保证表结构完整
        db.create_all()
    
    from waitress import serve
    print("🚀 UUID 离线同步架构版启动 (含旧数据迁移)...")
    serve(app, host='0.0.0.0', port=5000, threads=16, channel_timeout=300, cleanup_interval=300, outbuf_overflow=20 * 1024 * 1024, inbuf_overflow=20 * 1024 * 1024, connection_limit=200)